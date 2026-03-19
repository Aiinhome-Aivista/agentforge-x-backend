"""
File parsers: extract clean text / structured data from uploaded files.
Supports: PDF, DOCX, TXT, CSV, XLSX
"""

import os
import io
import logging
import chardet
import pandas as pd
import pdfplumber
from docx import Document
from typing import Tuple, Dict, Any

logger = logging.getLogger(__name__)


def detect_source_type(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower().lstrip(".")
    mapping = {
        "pdf": "pdf",
        "docx": "docx",
        "doc": "docx",
        "txt": "txt",
        "csv": "csv",
        "xlsx": "erp_dump",
        "xls": "erp_dump",
    }
    return mapping.get(ext, "txt")


def parse_file(file_bytes: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
    """
    Returns:
        text:     clean extracted text for LLM
        metadata: dict with extra structured info (tables, columns, shape etc.)
    """
    source_type = detect_source_type(filename)

    if source_type == "pdf":
        return _parse_pdf(file_bytes)
    elif source_type == "docx":
        return _parse_docx(file_bytes)
    elif source_type == "csv":
        return _parse_csv(file_bytes, filename)
    elif source_type == "erp_dump":
        return _parse_xlsx(file_bytes, filename)
    else:
        return _parse_txt(file_bytes)


# ── PDF ───────────────────────────────────────────────────────────────────────

def _parse_pdf(file_bytes: bytes) -> Tuple[str, Dict]:
    text_parts = []
    tables_found = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(f"[Page {i+1}]\n{page_text}")

            tables = page.extract_tables()
            for table in tables:
                if table:
                    df = pd.DataFrame(table[1:], columns=table[0])
                    tables_found.append(df.to_string(index=False))

    full_text = "\n\n".join(text_parts)
    if tables_found:
        full_text += "\n\n[TABLES EXTRACTED]\n" + "\n\n".join(tables_found)

    return full_text, {"pages": len(text_parts), "tables": len(tables_found)}


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _parse_docx(file_bytes: bytes) -> Tuple[str, Dict]:
    doc = Document(io.BytesIO(file_bytes))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            style = para.style.name if para.style else ""
            prefix = "# " if "Heading 1" in style else ("## " if "Heading" in style else "")
            parts.append(f"{prefix}{para.text}")

    table_texts = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        table_texts.append("\n".join(rows))

    full_text = "\n".join(parts)
    if table_texts:
        full_text += "\n\n[TABLES]\n" + "\n\n".join(table_texts)

    return full_text, {"paragraphs": len(parts), "tables": len(table_texts)}


# ── CSV (ERP Dump) ────────────────────────────────────────────────────────────

def _parse_csv(file_bytes: bytes, filename: str) -> Tuple[str, Dict]:
    detected = chardet.detect(file_bytes)
    encoding = detected.get("encoding", "utf-8") or "utf-8"

    try:
        df = pd.read_csv(io.BytesIO(file_bytes), encoding=encoding, nrows=500)
    except Exception:
        df = pd.read_csv(io.BytesIO(file_bytes), encoding="latin-1", nrows=500)

    summary = _dataframe_to_summary(df, filename)
    return summary, {
        "rows": len(df),
        "columns": list(df.columns),
        "shape": df.shape,
    }


# ── XLSX ──────────────────────────────────────────────────────────────────────

def _parse_xlsx(file_bytes: bytes, filename: str) -> Tuple[str, Dict]:
    xl = pd.ExcelFile(io.BytesIO(file_bytes))
    parts = []
    all_columns = {}

    for sheet_name in xl.sheet_names:
        df = xl.parse(sheet_name, nrows=300)
        parts.append(f"[Sheet: {sheet_name}]\n{_dataframe_to_summary(df, sheet_name)}")
        all_columns[sheet_name] = list(df.columns)

    return "\n\n".join(parts), {"sheets": xl.sheet_names, "columns_by_sheet": all_columns}


# ── TXT ───────────────────────────────────────────────────────────────────────

def _parse_txt(file_bytes: bytes) -> Tuple[str, Dict]:
    detected = chardet.detect(file_bytes)
    encoding = detected.get("encoding", "utf-8") or "utf-8"
    try:
        text = file_bytes.decode(encoding)
    except Exception:
        text = file_bytes.decode("latin-1")
    return text, {"chars": len(text)}


# ── Helpers ───────────────────────────────────────────────────────────────────

def _dataframe_to_summary(df: pd.DataFrame, name: str) -> str:
    """Convert a DataFrame into a text summary suitable for LLM input."""
    lines = [
        f"Dataset: {name}",
        f"Rows: {len(df)} | Columns: {len(df.columns)}",
        f"Columns: {', '.join(str(c) for c in df.columns)}",
        "",
        "Sample data (first 10 rows):",
        df.head(10).to_string(index=False),
        "",
        "Column statistics:",
    ]

    for col in df.columns:
        try:
            if df[col].dtype in ["int64", "float64"]:
                lines.append(f"  {col}: min={df[col].min()}, max={df[col].max()}, "
                              f"mean={df[col].mean():.2f}, nulls={df[col].isnull().sum()}")
            else:
                top_vals = df[col].value_counts().head(5).to_dict()
                lines.append(f"  {col}: unique={df[col].nunique()}, "
                              f"top_values={top_vals}, nulls={df[col].isnull().sum()}")
        except Exception:
            lines.append(f"  {col}: (unparseable)")

    return "\n".join(lines)
